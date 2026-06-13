#!/usr/bin/env python3
"""
Word ↔ Markdown 双向转换工具
========================================
用法：
  1. docx → md:  python convert_docx.py 文件.docx
  2. md → docx:   python convert_docx.py 文件.md
  3. 批量转换:     python convert_docx.py 文件1.docx 文件2.md ...
  4. 文件夹监听:   python convert_docx.py --watch [文件夹路径]
  5. 指定输出目录: python convert_docx.py 文件.docx -o ./output/

依赖安装: pip install -r requirements.txt
"""

import argparse
import os
import sys
import time
from pathlib import Path
from typing import Optional

# 修复 Windows 中文终端 GBK 编码问题
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# 过滤 C 库层面的 libpng iCCP 警告（无害，但干扰用户）
class _LibPngFilter:
    def __init__(self, stream):
        self.stream = stream
    def write(self, s):
        if "iCCP" not in s and "libpng warning" not in s:
            self.stream.write(s)
    def flush(self):
        self.stream.flush()
    def __getattr__(self, name):
        return getattr(self.stream, name)

sys.stderr = _LibPngFilter(sys.stderr)

# ============================================================
# 转换引擎选择
# ============================================================

def convert_with_mammoth(docx_path: str, md_path: str, extract_images: bool = False) -> bool:
    """使用 mammoth 库转换（推荐，纯 Python，安装简单）"""
    try:
        import mammoth
    except ImportError:
        return False

    # 图片处理
    style_map = ""
    if extract_images:
        img_dir = Path(md_path).parent / f"{Path(md_path).stem}_images"
        img_dir.mkdir(parents=True, exist_ok=True)
        from mammoth.documents import Image
        from mammoth.conversion import convert

        def convert_image(image):
            ext = image.content_type.split("/")[-1]
            img_path = img_dir / f"img_{hash(image.open)}_{len(os.listdir(img_dir))}.{ext}"
            with open(img_path, "wb") as f:
                f.write(image.open().read())
            return {"src": str(img_path.name)}

        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_markdown(f, convert_image=mammoth.images.img_element(convert_image))
    else:
        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_markdown(f)

    if result.messages:
        for msg in result.messages:
            print(f"  [!] {msg}")

    content = result.value

    # 后处理：清理多余空行
    import re
    content = re.sub(r"\n{4,}", "\n\n\n", content)

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content)

    return True


def convert_with_pandoc(docx_path: str, md_path: str, extract_images: bool = False) -> bool:
    """使用 pandoc 转换（需要安装 pandoc，转换质量最高）"""
    import subprocess
    import shutil

    if not shutil.which("pandoc"):
        return False

    cmd = [
        "pandoc",
        docx_path,
        "-f", "docx",
        "-t", "gfm",
        "--wrap=none",
        "--markdown-headings=atx",
        "-o", md_path,
    ]

    if extract_images:
        img_dir = Path(md_path).parent / f"{Path(md_path).stem}_images"
        img_dir.mkdir(parents=True, exist_ok=True)
        cmd.extend(["--extract-media", str(img_dir)])

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            print(f"   [!] pandoc 警告: {result.stderr.strip()}")
            return False  # 文件可能已创建但有问题
        return True
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return False


def _para_to_markdown(para) -> str:
    """将 python-docx 段落对象转为 markdown，保留格式"""
    from docx.oxml.ns import qn

    # 检测样式
    style_name = para.style.name.lower() if para.style else ""

    # ---- 标题检测 ----
    if style_name.startswith("heading"):
        try:
            level = int(style_name.replace("heading", "").strip().split()[0])
            level = min(level, 6)
        except (ValueError, IndexError):
            level = 1
        prefix = f"{'#' * level} "
        text = para.text.strip()
        return f"{prefix}{text}" if text else ""

    # ---- 列表检测 ----
    is_list = style_name.startswith("list") or any(
        run._element.rPr is not None and
        run._element.rPr.find(qn('w:rStyle')) is not None and
        'list' in (run._element.rPr.find(qn('w:rStyle')).get(qn('w:val')) or '').lower()
        for run in para.runs if hasattr(run, '_element')
    )

    # ---- 构建格式化文本 ----
    if not para.runs:
        text = para.text.strip()
        return f"- {text}" if is_list and text else text

    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue

        is_bold = run.bold
        is_italic = run.italic

        # 应用 markdown 格式
        if is_bold and is_italic:
            text = f"***{text}***"
        elif is_bold:
            text = f"**{text}**"
        elif is_italic:
            text = f"*{text}*"

        parts.append(text)

    result = "".join(parts).strip()

    if is_list and result:
        result = f"- {result}"

    return result


def _table_to_markdown(table) -> str:
    """将 python-docx 表格对象转为 markdown 表格"""
    if not table.rows:
        return ""

    lines = []
    # 表头
    headers = []
    for cell in table.rows[0].cells:
        text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
        headers.append(text)
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

    # 数据行
    for row in table.rows[1:]:
        cells = []
        for cell in row.cells:
            text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
            cells.append(text)
        lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)


def convert_with_python_docx(docx_path: str, md_path: str) -> bool:
    """
    使用 python-docx 增强转换 —— 按文档元素顺序处理，
    表格、段落、标题、列表混排时也能正确输出顺序。
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return False

    doc = Document(docx_path)

    # ---- 建立 XML 元素 → 对象映射 ----
    para_by_element = {}
    for para in doc.paragraphs:
        para_by_element[id(para._element)] = para

    table_by_element = {}
    for table in doc.tables:
        table_by_element[id(table._element)] = table

    # ---- 按文档顺序遍历 ----
    body = doc.element.body
    blocks = []  # 每个元素是一段 markdown 文本

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = para_by_element.get(id(child))
            if para:
                md = _para_to_markdown(para)
                if md:
                    blocks.append(md)
                else:
                    blocks.append("")  # 空行
        elif tag == "tbl":
            table = table_by_element.get(id(child))
            if table:
                blocks.append(_table_to_markdown(table))
        # 忽略其他元素（图片等）

    content = "\n\n".join(blocks)

    # 清理多余空行（但保留表格和代码块结构）
    import re
    content = re.sub(r"\n{4,}", "\n\n\n", content)

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content)

    return True


# ============================================================
# Markdown → Word (.docx) 转换
# ============================================================

def _parse_inline_formatting(text: str):
    """
    解析行内 markdown 格式，返回 (text_segments, formats) 的列表。
    formats: 每个 segment 的格式字典 {'bold': bool, 'italic': bool}
    """
    import re

    segments = []
    # 匹配: ***bold+italic***, **bold**, *italic*, __bold__, _italic_
    pattern = r'(\*\*\*|___|\*\*|__|\*|_)(.+?)\1'

    last_end = 0
    for match in re.finditer(pattern, text):
        # 前面的纯文本
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            if plain:
                segments.append((plain, {"bold": False, "italic": False}))

        marker = match.group(1)
        content = match.group(2)

        if marker in ("***", "___"):
            segments.append((content, {"bold": True, "italic": True}))
        elif marker in ("**", "__"):
            segments.append((content, {"bold": True, "italic": False}))
        elif marker in ("*", "_"):
            segments.append((content, {"bold": False, "italic": True}))

        last_end = match.end()

    # 剩余纯文本
    if last_end < len(text):
        plain = text[last_end:]
        if plain:
            segments.append((plain, {"bold": False, "italic": False}))

    # 如果没有匹配到任何格式标记，返回纯文本
    if not segments:
        segments.append((text, {"bold": False, "italic": False}))

    return segments


def _add_formatted_paragraph(doc, text: str, style: str = None):
    """向文档添加带格式的段落"""
    if not text.strip():
        p = doc.add_paragraph("")
        if style:
            p.style = doc.styles[style]
        return p

    segments = _parse_inline_formatting(text)

    if style and style.startswith("Heading"):
        # 标题不需要 inline 格式，直接设置文本
        level = int(style.replace("Heading ", ""))
        doc.add_heading(text.strip(), level=level)
        return

    p = doc.add_paragraph()
    for seg_text, fmt in segments:
        run = p.add_run(seg_text)
        if fmt["bold"]:
            run.bold = True
        if fmt["italic"]:
            run.italic = True

    return p


def _is_table_row(line: str) -> bool:
    """判断是否为 markdown 表格行"""
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _is_table_separator(line: str) -> bool:
    """判断是否为表格分隔行 (| --- | --- |)"""
    stripped = line.strip().replace(" ", "")
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return False
    # 分隔行只包含 |, -, : 字符
    inner = stripped[1:-1]
    return all(c in "|-:" for c in inner) and "-" in inner


def _parse_table_cells(line: str) -> list:
    """解析表格行，返回单元格文本列表"""
    stripped = line.strip()
    # 去掉首尾的 |
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    # 按 | 分割
    cells = [c.strip() for c in stripped.split("|")]
    return cells


def _is_ordered_list_item(line: str) -> tuple:
    """判断是否为有序列表项，返回 (是否, 序号, 内容)"""
    import re
    match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
    if match:
        indent = len(match.group(1))
        number = match.group(2)
        content = match.group(3)
        return (True, number, content, indent)
    return (False, "", "", 0)


def _is_unordered_list_item(line: str) -> tuple:
    """判断是否为无序列表项，返回 (是否, 内容, 缩进)"""
    import re
    match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
    if match and not line.strip().startswith("---"):
        indent = len(match.group(1))
        content = match.group(2)
        return (True, content, indent)
    return (False, "", 0)


def convert_md_to_docx(
    md_path: str,
    output_dir: Optional[str] = None,
) -> bool:
    """
    将 .md 文件转换为 .docx 文件。
    支持: 标题、段落、表格、有序/无序列表、粗体、斜体、粗斜体
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("[X] 需要安装 python-docx: pip install python-docx")
        return False

    md_file = Path(md_path)

    if not md_file.exists():
        print(f"[X] 文件不存在: {md_path}")
        return False

    if md_file.suffix.lower() not in (".md", ".markdown"):
        print(f"[X] 不支持的文件类型: {md_file.suffix}（仅支持 .md）")
        return False

    # 确定输出路径
    if output_dir:
        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
    else:
        out_dir = md_file.parent

    docx_file = out_dir / f"{md_file.stem}.docx"
    file_size = md_file.stat().st_size
    print(f">> 正在转换: {md_file.name} ({file_size / 1024:.1f} KB)")

    # 读取 markdown 内容
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # ---- 跳过空行 ----
        if not line:
            i += 1
            continue

        # ---- 表格检测 ----
        if _is_table_row(line):
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            # 过滤掉分隔行
            data_rows = [l for l in table_lines if not _is_table_separator(l)]
            if len(data_rows) < 1:
                continue

            # 解析表格
            header_cells = _parse_table_cells(data_rows[0])
            num_cols = len(header_cells)

            table = doc.add_table(rows=len(data_rows), cols=num_cols)
            table.style = 'Light Grid Accent 1'

            for row_idx, row_line in enumerate(data_rows):
                cells = _parse_table_cells(row_line)
                for col_idx, cell_text in enumerate(cells[:num_cols]):
                    cell = table.rows[row_idx].cells[col_idx]
                    # 清除默认段落
                    cell.paragraphs[0].clear()
                    # 解析并添加格式化文本
                    segments = _parse_inline_formatting(cell_text)
                    for seg_text, fmt in segments:
                        run = cell.paragraphs[0].add_run(seg_text)
                        if fmt["bold"]:
                            run.bold = True
                        if fmt["italic"]:
                            run.italic = True

            # 表格后添加空行
            doc.add_paragraph("")
            continue

        # ---- 标题检测 ----
        if line.startswith("#"):
            # 计算 # 的数量
            level = 0
            for ch in line:
                if ch == "#":
                    level += 1
                else:
                    break
            level = min(level, 6)
            heading_text = line[level:].strip()
            if heading_text:
                doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # ---- 有序列表 ----
        is_ol, num, content, indent = _is_ordered_list_item(line)
        if is_ol:
            list_items = []
            current_indent = indent
            while i < len(lines):
                is_ol_i, _, cont_i, ind_i = _is_ordered_list_item(lines[i].rstrip())
                if is_ol_i and ind_i == current_indent:
                    list_items.append(cont_i)
                    i += 1
                elif not lines[i].strip():
                    i += 1
                    # 检查下一行是否仍是列表
                    if i < len(lines):
                        is_next, _, _, _ = _is_ordered_list_item(lines[i].rstrip())
                        if not is_next:
                            break
                else:
                    break

            for idx, item_text in enumerate(list_items, 1):
                p = doc.add_paragraph()
                p.style = doc.styles['List Number']
                p.clear()
                segments = _parse_inline_formatting(item_text)
                for seg_text, fmt in segments:
                    run = p.add_run(seg_text)
                    if fmt["bold"]:
                        run.bold = True
                    if fmt["italic"]:
                        run.italic = True
            continue

        # ---- 无序列表 ----
        is_ul, content, indent = _is_unordered_list_item(line)
        if is_ul:
            list_items = []
            current_indent = indent
            while i < len(lines):
                is_ul_i, cont_i, ind_i = _is_unordered_list_item(lines[i].rstrip())
                if is_ul_i and ind_i == current_indent:
                    list_items.append(cont_i)
                    i += 1
                elif not lines[i].strip():
                    i += 1
                    if i < len(lines):
                        is_next, _, _ = _is_unordered_list_item(lines[i].rstrip())
                        if not is_next:
                            break
                else:
                    break

            for item_text in list_items:
                p = doc.add_paragraph()
                p.style = doc.styles['List Bullet']
                p.clear()
                segments = _parse_inline_formatting(item_text)
                for seg_text, fmt in segments:
                    run = p.add_run(seg_text)
                    if fmt["bold"]:
                        run.bold = True
                    if fmt["italic"]:
                        run.italic = True
            continue

        # ---- 水平线/分隔符 ----
        if line.strip() in ("---", "***", "___", "* * *", "- - -"):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # ---- 普通段落 ----
        _add_formatted_paragraph(doc, line)
        i += 1

    # 保存
    doc.save(str(docx_file))

    docx_size = docx_file.stat().st_size
    print(f"   [OK] 使用引擎: python-docx")
    print(f"   [->] 输出文件: {docx_file}")
    print(f"   [=] 大小: {docx_size / 1024:.1f} KB ({docx_size / max(file_size, 1):.1f}x 膨胀比)")

    return True


# ============================================================
# 核心转换函数
# ============================================================

def convert_docx_to_md(
    docx_path: str,
    output_dir: Optional[str] = None,
    extract_images: bool = False,
    force_engine: Optional[str] = None,
) -> bool:
    """
    将 .docx 文件转换为 .md 文件。
    自动选择最佳可用引擎：pandoc > mammoth > python-docx
    """
    docx_file = Path(docx_path)

    if not docx_file.exists():
        print(f"[X] 文件不存在: {docx_path}")
        return False

    if docx_file.suffix.lower() not in (".docx", ".doc"):
        print(f"[X] 不支持的文件类型: {docx_file.suffix}（仅支持 .docx）")
        return False

    # 确定输出路径
    if output_dir:
        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
    else:
        out_dir = docx_file.parent

    md_file = out_dir / f"{docx_file.stem}.md"

    file_size = docx_file.stat().st_size
    print(f">> 正在转换: {docx_file.name} ({file_size / 1024:.1f} KB)")

    # 按优先级尝试各引擎
    engines = []

    if force_engine:
        engines = [force_engine]
    else:
        engines = ["pandoc", "python-docx", "mammoth"]

    for engine in engines:
        if engine == "pandoc":
            if convert_with_pandoc(str(docx_file), str(md_file), extract_images):
                print(f"   [OK] 使用引擎: pandoc")
                break
        elif engine == "mammoth":
            if convert_with_mammoth(str(docx_file), str(md_file), extract_images):
                print(f"   [OK] 使用引擎: mammoth")
                break
        elif engine == "python-docx":
            if convert_with_python_docx(str(docx_file), str(md_file)):
                print(f"   [OK] 使用引擎: python-docx（含表格支持）")
                break
    else:
        print(f"[X] 没有可用的转换引擎！")
        print(f"   请至少安装以下之一：")
        print(f"   1. pip install mammoth     （推荐，纯 Python）")
        print(f"   2. pip install python-docx （基础支持）")
        print(f"   3. 安装 pandoc: https://pandoc.org/installing.html（质量最佳）")
        return False

    md_size = md_file.stat().st_size
    print(f"   [->] 输出文件: {md_file}")
    print(f"   [=] 大小: {md_size / 1024:.1f} KB ({file_size / max(md_size, 1):.1f}x 压缩比)")

    return True


# ============================================================
# 文件夹监听模式
# ============================================================

def watch_folder(
    folder_path: str,
    output_dir: Optional[str] = None,
    extract_images: bool = False,
    force_engine: Optional[str] = None,
    interval: int = 3,
    delete_after: bool = False,
):
    """
    监听文件夹，自动转换新放入的 .docx 文件。
    按 Ctrl+C 停止监听。
    """
    folder = Path(folder_path)
    if not folder.is_dir():
        print(f"[X] 不是有效文件夹: {folder_path}")
        return

    processed = set()

    print(f"[*] 正在监听文件夹: {folder}")
    print(f"   间隔: {interval} 秒 | 按 Ctrl+C 停止\n")

    try:
        while True:
            docx_files = list(folder.glob("*.docx")) + list(folder.glob("*.doc"))
            for f in docx_files:
                if f.name not in processed:
                    print(f"\n[NEW] 检测到新文件: {f.name}")
                    success = convert_docx_to_md(
                        str(f), output_dir=output_dir,
                        extract_images=extract_images,
                        force_engine=force_engine,
                    )
                    if success and delete_after:
                        f.unlink()
                        print(f"   [DEL] 已删除原文件: {f.name}")
                    processed.add(f.name)

            # 清理已删除文件的记录
            current_files = {f.name for f in docx_files}
            processed = processed & current_files

            time.sleep(interval)

    except KeyboardInterrupt:
        print("\n\n[END] 监听已停止。")


# ============================================================
# 命令行入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Word ↔ Markdown 双向转换工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # docx → md
  python convert_docx.py 起诉状.docx
  python convert_docx.py a.docx b.docx -o ./output/

  # md → docx
  python convert_docx.py 法律意见书.md
  python convert_docx.py 文件.md -o ./output/

  # 文件夹监听（支持 .docx 和 .md 互转）
  python convert_docx.py --watch ./待转换文件夹/
  python convert_docx.py --watch . --delete
        """,
    )

    parser.add_argument(
        "files", nargs="*",
        help="要转换的文件（.docx 或 .md，可多个，自动识别方向）"
    )
    parser.add_argument(
        "-o", "--output", dest="output_dir",
        help="输出目录（默认与原文件同目录）"
    )
    parser.add_argument(
        "-w", "--watch", dest="watch_folder",
        help="监听文件夹，自动转换新放入的 .docx 或 .md 文件"
    )
    parser.add_argument(
        "--delete", action="store_true",
        help="（监听模式）转换成功后删除原文件"
    )
    parser.add_argument(
        "--extract-images", action="store_true",
        help="（docx→md）提取文档中的图片到单独文件夹"
    )
    parser.add_argument(
        "--engine", choices=["pandoc", "mammoth", "python-docx"],
        help="（docx→md）强制使用指定转换引擎"
    )
    parser.add_argument(
        "--interval", type=int, default=3,
        help="（监听模式）扫描间隔，单位秒（默认 3）"
    )

    args = parser.parse_args()

    # 监听模式
    if args.watch_folder:
        watch_folder(
            args.watch_folder,
            output_dir=args.output_dir,
            extract_images=args.extract_images,
            force_engine=args.engine,
            interval=args.interval,
            delete_after=args.delete,
        )
        return

    # 批量转换模式
    if not args.files:
        parser.print_help()
        return

    success_count = 0
    for file_path in args.files:
        f = Path(file_path)
        suffix = f.suffix.lower()

        if suffix in (".docx", ".doc"):
            if convert_docx_to_md(
                file_path,
                output_dir=args.output_dir,
                extract_images=args.extract_images,
                force_engine=args.engine,
            ):
                success_count += 1
        elif suffix in (".md", ".markdown"):
            if convert_md_to_docx(
                file_path,
                output_dir=args.output_dir,
            ):
                success_count += 1
        else:
            print(f"[X] 不支持的文件类型: {suffix}（仅支持 .docx / .md）")
        print()

    print(f"[OK] 完成: {success_count}/{len(args.files)} 个文件转换成功")


if __name__ == "__main__":
    main()
